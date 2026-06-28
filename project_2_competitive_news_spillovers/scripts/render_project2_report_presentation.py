from __future__ import annotations

import subprocess
from pathlib import Path
from textwrap import fill

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from matplotlib.patches import FancyBboxPatch
from matplotlib.ticker import FuncFormatter, PercentFormatter


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
FIG_DIR = REPORT_DIR / "figures"
MD_PATH = REPORT_DIR / "project2_report_working.md"
DOCX_PATH = REPORT_DIR / "project2_report.docx"


def apply_base_style() -> None:
    plt.style.use("seaborn-v0_8-whitegrid")
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 10,
            "axes.titlesize": 13,
            "axes.labelsize": 10,
            "xtick.labelsize": 9,
            "ytick.labelsize": 9,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "axes.edgecolor": "#c9ced6",
            "grid.color": "#d9dee6",
            "grid.linewidth": 0.8,
        }
    )


def save_dataset_construction_figure() -> Path:
    output = FIG_DIR / "figure1_dataset_construction.png"
    stages = [
        ("Raw ticker-level news rows", 55094),
        ("Cleaned ticker-level rows", 52538),
        ("Unique articles", 38131),
        ("LLM-classified sample", 1250),
        ("Relevant classified articles", 930),
        ("Event-panel rows", 4544),
    ]
    raw_count = stages[0][1]

    fig, ax = plt.subplots(figsize=(10.5, 6.0), dpi=300)
    ax.set_xlim(0, 1)
    ax.set_ylim(-0.5, len(stages) - 0.5)
    ax.axis("off")

    left = 0.18
    max_width = 0.66
    bar_height = 0.62
    colors = ["#123c69", "#1d4e89", "#2d6aa3", "#4c85b5", "#6d9fc3", "#91b8d1"]

    for idx, (label, value) in enumerate(stages):
        y = len(stages) - 1 - idx
        width = max_width * (value / raw_count)
        box = FancyBboxPatch(
            (left, y - bar_height / 2),
            width,
            bar_height,
            boxstyle="round,pad=0.02,rounding_size=0.015",
            linewidth=0,
            facecolor=colors[idx],
        )
        ax.add_patch(box)
        if idx < len(stages) - 1:
            next_width = max_width * (stages[idx + 1][1] / raw_count)
            ax.plot(
                [left + width, left + next_width],
                [y - bar_height / 2, y - 1 + bar_height / 2],
                color="#9aa7b6",
                linewidth=1.4,
            )

        ax.text(0.01, y, label, va="center", ha="left", fontsize=10, color="#1a1a1a")
        ax.text(
            min(left + width + 0.02, 0.96),
            y,
            f"{value:,.0f}",
            va="center",
            ha="left",
            fontsize=10,
            color="#1a1a1a",
            fontweight="bold",
        )
        ax.text(
            0.96,
            y - 0.22,
            f"{value / raw_count:.1%} of raw",
            va="center",
            ha="right",
            fontsize=8.5,
            color="#5c6670",
        )

    ax.text(
        0.01,
        len(stages) - 0.1,
        "Dataset Construction",
        fontsize=13,
        fontweight="bold",
        color="#123c69",
        ha="left",
    )
    ax.text(
        0.01,
        len(stages) - 0.45,
        "Pipeline counts from raw ticker-linked news to the final competitor event panel",
        fontsize=9.5,
        color="#5c6670",
        ha="left",
    )
    fig.tight_layout()
    fig.savefig(output, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return output


def save_llm_distribution_figure() -> Path:
    output = FIG_DIR / "figure2_llm_relevance_distribution.png"
    df = pd.read_csv(ROOT / "outputs" / "tables" / "report_llm_label_distribution.csv")
    labels = [
        "Target company news",
        "Market roundup\nbut relevant",
        "Competitor company\nnews",
        "Industry news",
        "Macro-policy news",
        "Not relevant",
    ]
    values = df["article_count"].tolist()
    colors = ["#123c69", "#3b6ea5", "#5e91b3", "#80b5ad", "#a8cfa0", "#cfcfcf"]

    fig, ax = plt.subplots(figsize=(9.6, 5.4), dpi=300)
    bars = ax.bar(range(len(values)), values, color=colors, edgecolor="none", width=0.68)
    ax.set_xticks(range(len(values)), labels)
    ax.set_ylabel("Article count")
    ax.set_title("LLM Relevance-Type Distribution", pad=10, color="#123c69", fontweight="bold")
    ax.set_axisbelow(True)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    for bar, value in zip(bars, values):
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            bar.get_height() + 8,
            f"{value:,}",
            ha="center",
            va="bottom",
            fontsize=9,
            fontweight="bold",
            color="#1a1a1a",
        )

    fig.tight_layout()
    fig.savefig(output, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return output


def save_expected_effect_figure() -> Path:
    output = FIG_DIR / "figure3_mean_return_by_expected_effect.png"
    df = pd.read_csv(ROOT / "outputs" / "tables" / "mean_returns_by_label.csv")
    subset = df.loc[df["group_name"] == "expected_competitor_effect"].copy()
    ordered = [
        "same_direction_contagion",
        "opposite_direction_competition",
        "positive_for_competitors",
        "negative_for_competitors",
        "neutral_or_no_clear_effect",
    ]
    label_map = {
        "same_direction_contagion": "Same-direction\ncontagion",
        "opposite_direction_competition": "Opposite-direction\ncompetition",
        "positive_for_competitors": "Positive for\ncompetitors",
        "negative_for_competitors": "Negative for\ncompetitors",
        "neutral_or_no_clear_effect": "Neutral or\nno clear effect",
    }
    subset = subset.set_index("group_value").loc[ordered].reset_index()
    values = subset["competitor_abret_spy_t1_mean"].tolist()
    colors = ["#1d6996" if value >= 0 else "#c05640" for value in values]

    fig, ax = plt.subplots(figsize=(9.8, 5.6), dpi=300)
    bars = ax.bar(range(len(values)), values, color=colors, width=0.68, edgecolor="none")
    ax.axhline(0, color="#4f5b66", linewidth=1.1)
    ax.set_xticks(range(len(values)), [label_map[key] for key in ordered])
    ax.set_ylabel("Mean next-day abnormal return")
    ax.yaxis.set_major_formatter(PercentFormatter(xmax=1, decimals=2))
    ax.set_title(
        "Mean Competitor Abnormal Return by Expected Competitor Effect",
        pad=10,
        color="#123c69",
        fontweight="bold",
    )
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    for bar, value in zip(bars, values):
        va = "bottom" if value >= 0 else "top"
        offset = 0.00035 if value >= 0 else -0.00035
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            value + offset,
            f"{value:+.2%}",
            ha="center",
            va=va,
            fontsize=9,
            fontweight="bold",
            color="#1a1a1a",
        )

    fig.tight_layout()
    fig.savefig(output, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return output


def save_table5_image() -> Path:
    output = FIG_DIR / "table5_main_empirical_results.png"
    rows = [
        ["Broad", "Next-day abnormal return", "Industry news", "661", "0.0062", "0.0036", "Positive and statistically significant"],
        ["Strict", "Next-day abnormal return", "Industry news", "661", "0.0062", "0.0031", "Positive and statistically significant"],
        ["Broad", "CAR(1,3)", "Same-direction contagion", "3,093", "0.0113", "0.0381", "Positive in cumulative window"],
        ["Strict", "CAR(1,3)", "Same-direction contagion", "1,760", "0.0149", "0.0125", "Stronger in stricter sample"],
        ["Broad", "Next-day abnormal return", "Opposite-direction competition", "39", "0.0081", "0.3553", "Positive sign, but sparse and imprecise"],
        ["Strict", "Next-day abnormal return", "Opposite-direction competition", "27", "0.0072", "0.5442", "Positive sign, but sparse and imprecise"],
    ]
    columns = ["Sample", "Outcome", "Label", "Rows", "Coef.", "p-value", "Interpretation"]
    wrapped_rows = []
    for row in rows:
        wrapped_rows.append(
            [
                row[0],
                fill(row[1], width=18),
                fill(row[2], width=22),
                row[3],
                row[4],
                row[5],
                fill(row[6], width=28),
            ]
        )

    fig, ax = plt.subplots(figsize=(13.2, 3.8), dpi=300)
    ax.axis("off")
    table = ax.table(
        cellText=wrapped_rows,
        colLabels=columns,
        cellLoc="left",
        colLoc="left",
        bbox=[0, 0, 1, 1],
        colWidths=[0.09, 0.16, 0.19, 0.07, 0.08, 0.09, 0.32],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.5)

    for (row_idx, col_idx), cell in table.get_celld().items():
        cell.set_edgecolor("#bfc7d1")
        cell.set_linewidth(0.8)
        if row_idx == 0:
            cell.set_facecolor("#e9eef5")
            cell.set_text_props(weight="bold", color="#123c69")
        else:
            cell.set_facecolor("#f8fafc" if row_idx % 2 == 1 else "white")
            if col_idx in {3, 4, 5}:
                cell._loc = "center"

    fig.tight_layout()
    fig.savefig(output, dpi=300, bbox_inches="tight", pad_inches=0.02)
    plt.close(fig)
    return output


def render_base_docx() -> None:
    subprocess.run(
        ["pandoc", str(MD_PATH), "-o", str(DOCX_PATH)],
        check=True,
        cwd=ROOT,
    )


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._p
    props = p.pPr
    for child in list(p):
        p.remove(child)
    if props is not None:
        p.insert(0, props)


def set_caption(paragraph: Paragraph, text: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = "Calibri"


def insert_image_at_placeholder(
    doc: Document,
    placeholder: str,
    image_path: Path,
    caption: str,
    width_inches: float,
) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() != placeholder:
            continue
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        caption_para = insert_paragraph_after(paragraph)
        set_caption(caption_para, caption)
        return
    raise ValueError(f"Placeholder not found: {placeholder}")


def build_docx_with_figures() -> None:
    doc = Document(DOCX_PATH)
    insert_image_at_placeholder(
        doc,
        "[[FIGURE1_DATASET_CONSTRUCTION]]",
        FIG_DIR / "figure1_dataset_construction.png",
        "Figure 1. Dataset construction from raw news rows to the final competitor event panel.",
        6.2,
    )
    insert_image_at_placeholder(
        doc,
        "[[FIGURE2_LLM_RELEVANCE_DISTRIBUTION]]",
        FIG_DIR / "figure2_llm_relevance_distribution.png",
        "Figure 2. Distribution of relevance-type labels in the LLM-classified sample.",
        6.1,
    )
    insert_image_at_placeholder(
        doc,
        "[[TABLE5_MAIN_EMPIRICAL_RESULTS]]",
        FIG_DIR / "table5_main_empirical_results.png",
        "Table 5. Main empirical results.",
        6.35,
    )
    insert_image_at_placeholder(
        doc,
        "[[FIGURE3_MEAN_RETURN_BY_EXPECTED_EFFECT]]",
        FIG_DIR / "figure3_mean_return_by_expected_effect.png",
        "Figure 3. Mean next-day competitor abnormal return by expected competitor-effect label.",
        6.15,
    )
    doc.save(DOCX_PATH)


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    apply_base_style()
    save_dataset_construction_figure()
    save_llm_distribution_figure()
    save_expected_effect_figure()
    save_table5_image()
    render_base_docx()
    build_docx_with_figures()


if __name__ == "__main__":
    main()
